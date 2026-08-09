import io
import json
import copy
import streamlit as st
from docx import Document
import google.generativeai as genai

# Page Configuration
st.set_page_config(
    page_title="DocMind AI - Precision Template Transfer Engine",
    page_icon="📄",
    layout="centered"
)

st.title("📄 DocMind AI: Precision Template Transfer Engine")
st.write("Upload a reference `.docx` template, paste your text, and transfer styles with 100% precision!")

# ---------------------------------------------------------
# 1. AI STRUCTURAL PARSER
# ---------------------------------------------------------
def ai_intelligent_parse(raw_text, api_key):
    """
    Parses unstructured text into semantically labeled blocks via Gemini API.
    """
    genai.configure(api_key=api_key)
    
    candidate_models = [
        'gemini-2.0-flash',
        'gemini-1.5-flash',
        'gemini-1.5-pro'
    ]
    
    try:
        available = [m.name.replace("models/", "") for m in genai.list_models() if 'generateContent' in m.supported_generation_methods]
        flash_available = [m for m in available if 'flash' in m]
        if flash_available:
            candidate_models = flash_available + candidate_models
        elif available:
            candidate_models = available + candidate_models
    except Exception:
        pass

    prompt = f"""
    You are an expert document structure parser.
    Analyze the raw text and divide it into structured logical blocks.
    Classify each block into one of these exact types: 'title', 'heading1', 'heading2', or 'body'.

    Return ONLY a raw JSON list of objects with the keys "type" and "text". Do not include markdown fences or extra prose.

    Raw Text:
    {raw_text}
    """

    response = None
    last_error = None

    for model_name in candidate_models:
        try:
            model = genai.GenerativeModel(model_name)
            res = model.generate_content(prompt)
            if res and res.text:
                response = res
                break
        except Exception as e:
            last_error = e
            continue

    if not response or not response.text:
        raise RuntimeError(f"Could not connect to Gemini API: {str(last_error)}")

    clean_json = response.text.strip().lstrip("```json").rstrip("```").strip()
    return json.loads(clean_json)

# ---------------------------------------------------------
# 2. XML PARAGRAPH CLONING ENGINE (SAFE FROM KEYERRORS)
# ---------------------------------------------------------
class PreciseXMLEngine:
    def __init__(self, docx_stream):
        self.doc = Document(docx_stream)
        self.heading_elem = None
        self.body_elem = None
        self._find_exemplar_paragraphs()

    def _find_exemplar_paragraphs(self):
        """Finds representative Heading and Body paragraph XML elements in the template."""
        for p in self.doc.paragraphs:
            text = p.text.strip()
            if not text:
                continue
            
            # Check if paragraph looks like a heading (bold or short line)
            is_bold = any(run.bold for run in p.runs if run.bold is not None)
            
            if (is_bold or len(text) < 50) and self.heading_elem is None:
                self.heading_elem = copy.deepcopy(p._element)
            elif len(text) >= 50 and self.body_elem is None:
                self.body_elem = copy.deepcopy(p._element)

        # Fallback to whatever paragraphs exist in the doc
        if self.heading_elem is None and len(self.doc.paragraphs) > 0:
            self.heading_elem = copy.deepcopy(self.doc.paragraphs[0]._element)
        if self.body_elem is None and len(self.doc.paragraphs) > 0:
            self.body_elem = copy.deepcopy(self.doc.paragraphs[-1]._element)

    def _set_paragraph_text(self, p_element, new_text):
        """Clears existing runs in cloned XML and inserts the new text while keeping formatting."""
        # Find all run elements inside the paragraph
        runs = p_element.xpath('.//w:r')
        if runs:
            # Retain formatting of the first run
            first_run = runs[0]
            # Remove text elements inside first run and inject new text
            for t in first_run.xpath('.//w:t'):
                t.getparent().remove(t)
            
            # Create new text node
            from docx.oxml import OxmlElement
            from docx.oxml.ns import qn
            t_elem = OxmlElement('w:t')
            t_elem.text = new_text
            first_run.append(t_elem)

            # Remove all additional runs to avoid duplicate old text
            for extra_run in runs[1:]:
                extra_run.getparent().remove(extra_run)
        else:
            # Fallback if no runs exist
            p_element.text = new_text

    def generate_formatted_doc(self, structured_blocks):
        body_elem = self.doc._body._element

        # 1. Clear out original text paragraphs
        for p in list(self.doc.paragraphs):
            body_elem.remove(p._element)

        # 2. Duplicate exemplar XML nodes for each block
        for block in structured_blocks:
            b_type = block.get("type", "body").lower()
            text = block.get("text", "").strip()

            if not text:
                continue

            # Pick Heading or Body XML exemplar
            if b_type in ['title', 'heading1', 'heading2'] and self.heading_elem is not None:
                new_p_elem = copy.deepcopy(self.heading_elem)
            elif self.body_elem is not None:
                new_p_elem = copy.deepcopy(self.body_elem)
            else:
                p = self.doc.add_paragraph(text)
                continue

            # Replace text inside XML element
            self._set_paragraph_text(new_p_elem, text)

            # Append XML paragraph directly to document body
            body_elem.append(new_p_elem)

        output_stream = io.BytesIO()
        self.doc.save(output_stream)
        output_stream.seek(0)
        return output_stream

# ---------------------------------------------------------
# 3. STREAMLIT UI
# ---------------------------------------------------------

api_key = st.secrets.get("GEMINI_API_KEY", "")

st.subheader("1. Upload Reference Template (.docx)")
uploaded_file = st.file_uploader("Upload sample Word file (Must be .docx)", type=["docx"])

st.subheader("2. Paste Content")
raw_user_input = st.text_area("Paste text or paragraphs here:", value="", height=280)

if st.button("🚀 Intelligently Parse & Format Document"):
    if not api_key:
        st.error("⚠️ GEMINI_API_KEY secret not found! Configure secrets in Streamlit Cloud settings.")
    elif uploaded_file is None:
        st.error("Please upload a `.docx` template file first!")
    elif not raw_user_input.strip():
        st.error("Please paste or type content into the box!")
    else:
        try:
            with st.spinner("🤖 Classifying content structure..."):
                structured_blocks = ai_intelligent_parse(raw_user_input, api_key)
                
            with st.spinner("🎨 Cloning template XML nodes for exact formatting..."):
                engine = PreciseXMLEngine(uploaded_file)
                output_doc_stream = engine.generate_formatted_doc(structured_blocks)

                st.success("Successfully generated accurately formatted document!")
                st.download_button(
                    label="📥 Download Formatted Word Document",
                    data=output_doc_stream,
                    file_name="DocMind_Formatted_Output.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
        except Exception as e:
            st.error(f"Error during parsing or document formatting: {str(e)}")
